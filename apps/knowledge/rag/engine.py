"""
Document parsing, chunking, embedding, vector storage.
Zero heavy dependencies: numpy + custom splitter + cosine similarity.
"""

import json, re, shutil
from pathlib import Path
import numpy as np
from django.conf import settings
from openai import OpenAI


def get_embedding_client() -> OpenAI:
    return OpenAI(
        api_key=settings.SILICONFLOW_API_KEY,
        base_url=settings.SILICONFLOW_BASE_URL,
    )


# ── Vector store ─────────────────────────────────────────────────

class VectorStore:
    def __init__(self, kb_id: str):
        self.store_dir = Path(settings.CHROMA_PERSIST_DIR) / f"kb_{kb_id}"
        self.store_dir.mkdir(parents=True, exist_ok=True)

    def _load(self):
        cf = self.store_dir / "chunks.json"
        ef = self.store_dir / "embeddings.npy"
        if not cf.exists() or not ef.exists():
            return [], None
        return json.loads(cf.read_text(encoding="utf-8")), np.load(ef)

    def _save(self, chunks, emb):
        (self.store_dir / "chunks.json").write_text(json.dumps(chunks, ensure_ascii=False), encoding="utf-8")
        np.save(self.store_dir / "embeddings.npy", emb)

    def add(self, doc_id, filename, chunks, embeddings):
        all_c, all_e = self._load()
        for i, (c, e) in enumerate(zip(chunks, embeddings)):
            all_c.append({"doc_id": doc_id, "filename": filename, "chunk_index": i, "content": c})
        new_e = np.array(embeddings) if all_e is None else np.vstack([all_e, embeddings])
        self._save(all_c, new_e)

    def remove_doc(self, doc_id):
        all_c, all_e = self._load()
        keep = [i for i, c in enumerate(all_c) if c["doc_id"] != doc_id]
        if not keep:
            shutil.rmtree(self.store_dir, ignore_errors=True)
            return
        self._save([all_c[i] for i in keep], all_e[keep])

    def query(self, q_emb, top_k=5):
        all_c, all_e = self._load()
        if all_e is None or len(all_c) == 0:
            return []
        q = np.array(q_emb)
        sims = np.dot(all_e, q) / (np.linalg.norm(all_e, axis=1) * np.linalg.norm(q) + 1e-10)
        top = np.argsort(sims)[::-1][:top_k]
        return [
            {"content": all_c[i]["content"][:500], "filename": all_c[i]["filename"], "chunk_index": all_c[i]["chunk_index"]}
            for i in top if sims[i] > 0.1
        ]

    def count(self):
        c, _ = self._load()
        return len(c) if c else 0


def delete_collection(kb_id):
    shutil.rmtree(Path(settings.CHROMA_PERSIST_DIR) / f"kb_{kb_id}", ignore_errors=True)

def get_store(kb_id):
    return VectorStore(kb_id)


# ── Parsing ──────────────────────────────────────────────────────

import io

def _read_with_fallback_encodings(fp):
    """Try utf-8 → gbk → latin-1 for text files."""
    for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            return fp.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return fp.read_text(encoding='utf-8', errors='replace')

def parse_pdf(fp):
    from PyPDF2 import PdfReader
    parts = []
    for page in PdfReader(str(fp)).pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n\n".join(parts)

def parse_docx(fp):
    from docx import Document
    doc = Document(str(fp))
    parts = []
    # 段落
    for p in doc.paragraphs:
        if p.text.strip():
            # 检测标题样式
            if p.style and p.style.name and p.style.name.startswith('Heading'):
                level = p.style.name.replace('Heading ', '').strip()
                prefix = '#' * min(int(level), 4) if level.isdigit() else ''
                parts.append(f"\n{prefix} {p.text.strip()}\n")
            else:
                parts.append(p.text.strip())
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)

def parse_txt(fp):
    return _read_with_fallback_encodings(fp)

PARSERS = {"pdf": parse_pdf, "docx": parse_docx, "doc": parse_docx, "txt": parse_txt, "md": parse_txt}

def parse_document(fp, ft):
    return PARSERS.get(ft.lower(), parse_txt)(fp)


# ── Chunking ─────────────────────────────────────────────────────

# 中文句子边界：。！？；：\n
# 英文句子边界：.!?;:\n
# 混合边界
_SENT_END = r'(?<=[。！？；：.!?;:\n])\s*'

def chunk_text(text, chunk_size=600, overlap=80):
    """Split text into overlapping chunks, preserving structure.

    Strategy: split by double-newline (paragraphs/sections) first,
    then split long paragraphs by sentence boundary, then force-split
    by character if still too long.
    """
    if not text or not text.strip():
        return []

    text = re.sub(r'\n{3,}', '\n\n', text)  # normalize
    text = re.sub(r' +', ' ', text)          # collapse spaces

    # 保留可能的 Markdown 标题 (#, ##, ###)
    sections = re.split(r'(\n(?=#{1,4} )|\n{2,})', text)

    chunks = []
    buf = ""

    for sec in sections:
        if not sec or sec.startswith('\n'):
            if buf:
                chunks.append(buf.strip())
                buf = ""
            continue

        sec = sec.strip()
        if not sec:
            continue

        # 短段落直接拼到缓冲区
        if len(buf) + len(sec) + 2 <= chunk_size:
            buf = (buf + "\n\n" + sec).strip()
        else:
            if buf:
                chunks.append(buf.strip())
            # 单段太长 → 按句子切
            if len(sec) > chunk_size:
                sents = re.split(_SENT_END, sec)
                sbuf = ""
                for s in sents:
                    s = s.strip()
                    if not s:
                        continue
                    if len(sbuf) + len(s) + 1 <= chunk_size:
                        sbuf = (sbuf + " " + s).strip() if sbuf else s
                    else:
                        if sbuf:
                            chunks.append(sbuf)
                        # 单句太长 → 按字符 force-split
                        if len(s) > chunk_size:
                            step = chunk_size - overlap
                            for i in range(0, len(s), step):
                                chunks.append(s[i:i+chunk_size])
                        else:
                            sbuf = s
                buf = sbuf if sbuf and len(sbuf) <= chunk_size else ""
            else:
                buf = sec

    if buf and buf.strip():
        chunks.append(buf.strip())

    # 过滤太短的片段
    return [c for c in chunks if len(c) >= 20]


# ── Embedding ────────────────────────────────────────────────────

def embed_chunks(chunks):
    client = get_embedding_client()
    resp = client.embeddings.create(model=settings.EMBEDDING_MODEL, input=chunks)
    return [d.embedding for d in resp.data]


# ── Retrieve ─────────────────────────────────────────────────────

def retrieve(kb_id, query, top_k=5):
    store = get_store(kb_id)
    if store.count() == 0:
        return []
    resp = get_embedding_client().embeddings.create(model=settings.EMBEDDING_MODEL, input=[query])
    return store.query(resp.data[0].embedding, top_k)


# ── Async task ───────────────────────────────────────────────────

def process_document_task(doc_id):
    from apps.knowledge.models import Document
    try:
        doc = Document.objects.get(id=doc_id)
    except Document.DoesNotExist:
        return
    doc.status = Document.Status.PROCESSING
    doc.save(update_fields=["status"])
    try:
        fp = Path(doc.file.path)
        content = parse_document(fp, doc.file_type)
        doc.content_text = content[:50000] if content else ""  # 保存前 50k 字符
        doc.save(update_fields=["content_text"])

        chunks = chunk_text(content)
        if not chunks:
            chunks = [content[:600]]

        # 分批 embedding（每次最多 50 个片段）
        all_embs = []
        batch_size = 50
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            all_embs.extend(embed_chunks(batch))

        get_store(str(doc.kb_id)).add(str(doc.id), doc.filename, chunks, all_embs)
        doc.chunk_count = len(chunks)
        doc.status = Document.Status.READY
        doc.save(update_fields=["chunk_count", "status"])
    except Exception as e:
        doc.status = Document.Status.ERROR
        doc.error_message = str(e)
        doc.save(update_fields=["status", "error_message"])

def remove_document_chunks(kb_id, doc_id):
    get_store(kb_id).remove_doc(doc_id)
